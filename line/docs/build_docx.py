# -*- coding: utf-8 -*-
"""Markdown を Word(.docx) に変換するスクリプト（引数で入力ファイルを指定可）。"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1] if len(sys.argv) > 1 else "TALK_ユーザーマニュアル.md"  # 入力Markdown
OUT = SRC.rsplit(".", 1)[0] + ".docx"       # 出力Word（拡張子だけ差し替え）

GREEN = RGBColor(0x4C, 0xAF, 0x50)          # アプリのメインカラー（緑）
GRAY = RGBColor(0x66, 0x66, 0x66)           # 補足テキストのグレー

doc = Document()

# 既定フォント（日本語対応）
style = doc.styles["Normal"]
style.font.name = "游ゴシック"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")


def set_cell_bg(cell, color_hex):
    """セルの背景色を設定する。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_runs_with_bold(paragraph, text):
    """**太字** 記法を解釈してランを追加する。"""
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if part == "":
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:                       # 奇数番目は太字部分
            run.bold = True


def parse_table(lines, start):
    """Markdownの表を解析して開始位置と行データを返す。"""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        cells = [c.strip() for c in line.strip("|").split("|")]
        if not re.match(r"^[\s:|-]+$", line.replace("|", "")):  # 区切り行は除外
            rows.append(cells)
        i += 1
    return rows, i


def add_table(rows):
    """解析済みの行データからWordの表を追加する。"""
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.paragraphs[0].text = ""
            add_runs_with_bold(cell.paragraphs[0], text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
                run.font.name = "游ゴシック"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
            if r_idx == 0:                   # 見出し行は緑背景＋白文字
                set_cell_bg(cell, "4CAF50")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
in_screenshot = False                        # スクリーンショット枠の中か
screenshot_buf = []

while i < len(lines):
    raw = lines[i]
    line = raw.rstrip()
    stripped = line.strip()

    # スクリーンショット指示ブロックをまとめて枠に入れる
    if stripped.startswith("【スクリーンショット】"):
        in_screenshot = True
        screenshot_buf = ["【スクリーンショット】"]
        i += 1
        while i < len(lines) and lines[i].strip() != "" and not lines[i].startswith("#") and not lines[i].startswith(">"):
            screenshot_buf.append(lines[i].rstrip())
            i += 1
        # 枠付きの段落として出力
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "dashed")
            e.set(qn("w:sz"), "6")
            e.set(qn("w:space"), "6")
            e.set(qn("w:color"), "4CAF50")
            pBdr.append(e)
        pPr.append(pBdr)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F1F8F1")
        pPr.append(shd)
        for j, bl in enumerate(screenshot_buf):
            run = p.add_run(bl)
            run.font.size = Pt(9)
            run.font.name = "游ゴシック"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
            if j == 0:
                run.bold = True
                run.font.color.rgb = GREEN
            if j != len(screenshot_buf) - 1:
                run.add_break()
        in_screenshot = False
        continue

    # 水平線 → 区切りとして無視（前段落に余白）
    if stripped == "---":
        i += 1
        continue

    # 見出し
    if stripped.startswith("# "):
        h = doc.add_heading(stripped[2:], level=1)
        for run in h.runs:
            run.font.color.rgb = GREEN
            run.font.name = "游ゴシック"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        i += 1
        continue
    if stripped.startswith("## "):
        h = doc.add_heading(stripped[3:], level=2)
        for run in h.runs:
            run.font.color.rgb = GREEN
            run.font.name = "游ゴシック"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        i += 1
        continue
    if stripped.startswith("### "):
        h = doc.add_heading(stripped[4:], level=3)
        for run in h.runs:
            run.font.name = "游ゴシック"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        i += 1
        continue

    # 表
    if stripped.startswith("|"):
        rows, ni = parse_table(lines, i)
        add_table(rows)
        i = ni
        continue

    # 引用（ワンポイント等）→ 薄い枠の段落
    if stripped.startswith(">"):
        quote_lines = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            quote_lines.append(lines[i].strip().lstrip(">").strip())
            i += 1
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "FFF8E1")      # 薄い黄色（アドバイス背景）
        pPr.append(shd)
        pBdr = OxmlElement("w:pBdr")
        e = OxmlElement("w:left")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "18")
        e.set(qn("w:space"), "8")
        e.set(qn("w:color"), "F5A623")
        pBdr.append(e)
        pPr.append(pBdr)
        for j, ql in enumerate(quote_lines):
            if ql == "":
                continue
            add_runs_with_bold(p, ql)
            if j != len(quote_lines) - 1:
                p.add_run().add_break()
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.name = "游ゴシック"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        continue

    # 箇条書き（- / チェックリスト）
    if stripped.startswith("- "):
        text = stripped[2:]
        if text.startswith("[ ] "):
            text = "☐ " + text[4:]
        p = doc.add_paragraph(style="List Bullet")
        add_runs_with_bold(p, text)
        for run in p.runs:
            run.font.name = "游ゴシック"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        i += 1
        continue

    # 番号付きリスト
    m = re.match(r"^\d+\.\s+(.*)", stripped)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_runs_with_bold(p, m.group(1))
        for run in p.runs:
            run.font.name = "游ゴシック"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        i += 1
        continue

    # 空行
    if stripped == "":
        i += 1
        continue

    # 通常の段落（太字・斜体を簡易処理）
    text = stripped
    if text.startswith("*") and text.endswith("*") and not text.startswith("**"):
        p = doc.add_paragraph()
        run = p.add_run(text.strip("*"))
        run.italic = True
        run.font.color.rgb = GRAY
        run.font.size = Pt(9)
    else:
        p = doc.add_paragraph()
        add_runs_with_bold(p, text)
    for run in p.runs:
        run.font.name = "游ゴシック"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
    i += 1

doc.save(OUT)
print("saved:", OUT)
